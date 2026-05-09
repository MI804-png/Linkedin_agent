"""
Extract all text from AutoApply_Thesis_v23.docx and divide into two parts
"""
from docx import Document
import os

# Load the generated thesis
thesis_path = r'd:\cv_portofolio\thesis\AutoApply_Thesis_v23.docx'

if not os.path.exists(thesis_path):
    print(f"ERROR: Thesis file not found at {thesis_path}")
    print("Regenerating thesis first...")
    os.system(r'd:\cv_portofolio\.venv\Scripts\python.exe d:\cv_portofolio\thesis\generate_thesis_v2.py')

doc = Document(thesis_path)

# Extract all text from all paragraphs
all_text = []
for para in doc.paragraphs:
    if para.text.strip():  # Only non-empty paragraphs
        all_text.append(para.text)

# Also extract text from tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.text.strip():
                    all_text.append(para.text)

# Combine all text
full_text = '\n\n'.join(all_text)

# Divide into two parts
mid_point = len(full_text) // 2
part1 = full_text[:mid_point]
part2 = full_text[mid_point:]

# Save to files
output_dir = r'd:\cv_portofolio\thesis'

part1_path = os.path.join(output_dir, 'thesis_text_part1.txt')
part2_path = os.path.join(output_dir, 'thesis_text_part2.txt')

with open(part1_path, 'w', encoding='utf-8') as f:
    f.write(part1)

with open(part2_path, 'w', encoding='utf-8') as f:
    f.write(part2)

# Print statistics
print(f"✓ Thesis extracted and divided into two parts:\n")
print(f"Total text length: {len(full_text):,} characters")
print(f"Part 1: {len(part1):,} characters")
print(f"Part 2: {len(part2):,} characters")
print(f"\nFiles saved:")
print(f"  • {part1_path}")
print(f"  • {part2_path}")

# Print first 500 chars of each part as preview
print(f"\n{'='*80}")
print(f"PART 1 PREVIEW (first 500 characters):")
print(f"{'='*80}")
print(part1[:500] + "...\n")

print(f"{'='*80}")
print(f"PART 2 PREVIEW (first 500 characters):")
print(f"{'='*80}")
print(part2[:500] + "...\n")
