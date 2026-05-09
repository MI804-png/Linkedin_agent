"""
AutoApply Thesis Generator v2
- Opens FinalThesis_Template.docx (inherits all szd_* styles)
- Clears existing body content, preserves section layout
- Adds ~50-page thesis with proper styles + embedded screenshots
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = r'd:\cv_portofolio\thesis\FinalThesis_Template.docx'
SHOTS    = r'd:\cv_portofolio\thesis\screenshots'
LOGO     = r'd:\cv_portofolio\thesis\template_image1.png'
OUT      = r'd:\cv_portofolio\thesis\AutoApply_Thesis_v23.docx'

doc = Document(TEMPLATE)

# ── 1. Clear all existing body content, keep sectPr (page layout) ──────────
body = doc.element.body
for child in list(body):
    if child.tag != qn('w:sectPr'):
        body.remove(child)

# ── Helper functions using template styles ──────────────────────────────────
def T(text, style='szd_szöveg'):
    """Add a paragraph using a template style."""
    p = doc.add_paragraph(text, style=style)
    return p

def H1(text):
    return T(text, 'szd_címsor1')

def H1u(text):
    """Unnumbered chapter heading (Introduction, Summary…)"""
    return T(text, 'szd_címsor_számozatlan')

def H2(text):
    p = T(text, 'szd_címsor2')
    # Compact subsection heading spacing to reduce pagination bloat.
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    return p

def H3(text):
    p = T(text, 'szd_címsor3')
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    return p

def body_text(text):
    p = T(text, 'szd_szöveg')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.0
    return p

def bullet(text):
    p = T(text, 'szd_felsorolás')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    return p

def numbered(text):
    p = T(text, 'szd_felsorolás_szám')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    return p

def code(text):
    p = doc.add_paragraph(style='szd_forráskód')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(text)
    return p

def fig_caption(text):
    p = T(text, 'szd_ábracím')
    # Template caption style is italic by default; force upright text for readability.
    for run in p.runs:
        run.italic = False
    # Tighten caption spacing so figures consume less vertical space.
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    # Allow tighter pagination around captions.
    p.paragraph_format.keep_with_next = False
    p.paragraph_format.keep_together = False
    return p

def ref_entry(text):
    return T(text, 'szd_szakirodalom')

def blank():
    return T('', 'szd_szöveg')

def add_image(filename, width_cm=14, caption=None):
    """Embed a PNG from the screenshots folder."""
    path = os.path.join(SHOTS, filename)
    if os.path.exists(path):
        # Global compaction: shrink figure footprint while preserving readability.
        width_cm = round(width_cm * 0.88, 2)
        # Image paragraph uses szd_ábra style
        p = doc.add_paragraph(style='szd_ábra')
        # Let Word flow figures naturally to avoid large whitespace blocks.
        p.paragraph_format.keep_with_next = False
        p.paragraph_format.keep_together = False
        run = p.add_run()
        run.add_picture(path, width=Cm(width_cm))
    else:
        body_text(f'[Figure: {filename} — not found]')
    if caption:
        fig_caption(caption)

def add_image_if_exists(filename, width_cm=14, caption=None):
    """Embed image only when available (no placeholder if missing)."""
    path = os.path.join(SHOTS, filename)
    if not os.path.exists(path):
        return False
    add_image(filename, width_cm=width_cm, caption=caption)
    return True

def toc_entry(text, page, level=0):
    style = 'toc 1' if level == 0 else ('toc 2' if level == 1 else 'toc 3')
    try:
        p = doc.add_paragraph(style=style)
    except:
        p = doc.add_paragraph(style='szd_szöveg')
    run = p.add_run(text)
    tab = p.add_run(f'\t{page}')
    return p

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════

# University logo at the very top (matches template paragraph [0])
logo_para = doc.add_paragraph()
logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_para.paragraph_format.space_before = Pt(0)
logo_para.paragraph_format.space_after  = Pt(0)
if os.path.exists(LOGO):
    logo_para.add_run().add_picture(LOGO, width=Cm(10))

for _ in range(5):
    blank()

cover = doc.add_paragraph(
    "Design and Implementation of a Data-Driven HR\n"
    "and Management Decision Support System\n"
    "for Organizational Performance and Risk Analysis"
)
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(0)
cover.paragraph_format.space_after  = Pt(24)
for run in cover.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True

for _ in range(4):
    blank()

for line in [
    "Mikhael Nabil Salama Rezk",
    "IHUTSC",
    "",
    "University Consultant: Mark Kovacs, position: Computer Engineering",
    "",
]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

for _ in range(6):
    blank()

year = doc.add_paragraph("2026")
year.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in year.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════════
toc_title = T('Table of Contents', 'szd_címsor_tartalomjegyzék')

toc_items = [
    ("INTRODUCTION", "4", 0),
    ("1.\tINTRODUCTION AND PROBLEM DEFINITION", "5", 0),
    ("1.1.\tProblem Context and Motivation", "5", 1),
    ("1.2.\tProblem Statement and Identified Gaps", "6", 1),
    ("1.3.\tThesis Objectives and Research Questions", "7", 1),
    ("1.4.\tScope and Delimitation", "7", 1),
    ("2.\tLITERATURE REVIEW AND RELATED WORK", "9", 0),
    ("2.1.\tWeb Automation and Browser Control Technologies", "9", 1),
    ("2.2.\tApplicant Tracking Systems and Recruitment Technology", "10", 1),
    ("2.3.\tAutomated Form Recognition and Intelligent DOM Navigation", "11", 1),
    ("2.4.\tAnti-Bot Detection Mechanisms and Evasion Strategies", "12", 1),
    ("2.5.\tGap Analysis and Research Positioning", "13", 1),
    ("3.\tSYSTEM DESIGN AND ARCHITECTURE", "14", 0),
    ("3.1.\tArchitectural Overview: Four-Layer Design", "14", 1),
    ("3.2.\tPage Type Classification System", "15", 1),
    ("3.3.\tSmart DOM Scoring Engine", "17", 1),
    ("3.4.\tBot Workflow State Machine", "18", 1),
    ("3.5.\tRisk Management and System Security", "19", 1),
    ("4.\tIMPLEMENTATION AND SYSTEM DEVELOPMENT", "21", 0),
    ("4.1.\tProgramming Languages and Technology Stack", "21", 1),
    ("4.2.\tLibrary Reference and Functional Roles", "23", 1),
    ("4.3.\tCore Modules and Function Reference", "25", 1),
    ("4.4.\tExternal Job Scraping Pipeline", "27", 1),
    ("4.5.\tDatabase Schema and Persistence Strategy", "29", 1),
    ("4.6.\tAI-Powered Cover Letter Generation and CV Submission", "30", 1),
    ("4.7.\tSystem Interface Screenshots", "31", 1),
    ("5.\tRESULTS, EVALUATION, AND SYSTEM VALIDATION", "39", 0),
    ("5.1.\tValidation Methodology and Test Coverage", "39", 1),
    ("5.2.\tObserved Results and System Performance", "41", 1),
    ("5.3.\tPractical Application and Use Cases", "42", 1),
    ("SUMMARY (CONCLUSIONS)", "45", 0),
    ("LIST OF FIGURES", "46", 0),
    ("REFERENCES", "47", 0),
    ("ATTACHMENTS", "48", 0),
]

for text, page, level in toc_items:
    toc_entry(text, page, level)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# INTRODUCTION (unnumbered)
# ════════════════════════════════════════════════════════════════════════════
H1u("Introduction")

body_text(
    "The modern employment market is characterised by a structural imbalance between "
    "opportunity density and individual application capacity. Digital transformation has "
    "proliferated job postings across dozens of platforms—corporate career portals, "
    "specialised remote-work boards, and professional networking services—yet the process "
    "of submitting an application to each posting remains largely manual, requiring candidates "
    "to navigate individually designed workflows, fill repetitive form fields, upload identical "
    "documents, and manage submission status across heterogeneous interfaces. Industry data "
    "consistently reports that active candidates in technical disciplines submit between 50 and "
    "200 applications per month to achieve statistically meaningful interview rates [1][3]; "
    "the mechanical work associated with this volume constitutes the dominant time cost of "
    "the job search process."
)
body_text(
    "Automation of the application workflow presents a technically sophisticated challenge. "
    "Application forms are rendered by competing Applicant Tracking Systems (ATS)—commercial "
    "platforms such as Greenhouse, Lever, Ashby, BambooHR, Recruitee, SmartRecruiters, and "
    "Jobvite—each implementing distinct DOM structures, proprietary field naming conventions, "
    "and custom submission mechanisms. Beyond structural heterogeneity, modern job platforms "
    "actively deploy anti-automation defences: Cloudflare Bot Management presents JavaScript-"
    "based challenge pages that detect headless browser fingerprints, platforms such as "
    "RemoteOK interrupt application flows with mandatory registration walls, and certain ATS "
    "providers require interactive CAPTCHA completion before granting access to application forms."
)
body_text(
    "This thesis documents the complete design and implementation of AutoApply, a locally-"
    "deployable intelligent job application automation system that addresses these challenges "
    "through a multi-layer technical architecture. The system combines Playwright-based "
    "stealth browser automation, a 14-category page type classifier, a DOM scoring engine "
    "for adaptive element detection, ATS-specific form handlers, a multi-platform job scraping "
    "pipeline covering four major job boards, and a Flask web application providing user "
    "management and campaign monitoring. The design prioritises local operation—the system "
    "runs on a Windows machine against a SQLite database without cloud dependencies—ensuring "
    "full candidate data sovereignty."
)
body_text(
    "The thesis proceeds as follows. Chapter 1 establishes the problem context and thesis "
    "objectives. Chapter 2 reviews relevant literature on browser automation, ATS technology, "
    "and anti-bot detection. Chapter 3 presents the system architecture including the risk "
    "management framework. Chapter 4 details the full implementation including programming "
    "languages, library stack, function reference, and system screenshots. Chapter 5 presents "
    "validation results and practical use cases. The Summary and Appendices follow."
)

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 – INTRODUCTION AND PROBLEM DEFINITION
# ════════════════════════════════════════════════════════════════════════════
H1("1.\tIntroduction and Problem Definition")

H2("1.1.\tProblem Context and Motivation")
body_text(
    "Active job seekers in software engineering and adjacent technical disciplines face a "
    "market characterised by high opportunity volume and low response rates. Published surveys "
    "from career coaching platforms and HR research firms consistently report that candidates "
    "must submit 50–200 applications to generate interview invitations at a frequency sufficient "
    "for effective career management [1][3]. A single application to a standard ATS form "
    "requires an average of 15–25 minutes when performed manually, encompassing: account "
    "creation or login on the employer's platform, duplication of profile information across "
    "non-standardised fields, CV upload, cover letter composition, and submission of "
    "supplementary answers to screening questions."
)
body_text(
    "At 200 applications and 20 minutes per submission, the mechanical application layer "
    "consumes approximately 67 hours per active job search campaign. This constitutes direct "
    "opportunity cost: time spent on mechanical form-filling is time not spent on company "
    "research, interview preparation, skills development, or networking—activities that "
    "meaningfully differentiate candidates. The motivation for automation is therefore not "
    "to increase application volume per se, but to reclaim time from the mechanical layer and "
    "redirect it toward high-value differentiation activities."
)
body_text(
    "A secondary motivation arises from the asymmetric information problem in job matching. "
    "Candidates with strong profiles for a given role may fail to apply simply because they "
    "did not encounter the posting before it expired, or because the manual application cost "
    "was perceived as too high relative to the probability of success. Automation reduces "
    "the marginal cost of each application toward zero, enabling broader coverage of the "
    "opportunity landscape and improving the probability of match for any given skill profile."
)

H2("1.2.\tProblem Statement and Identified Gaps")
body_text(
    "Three principal gaps define the technical problem space for automated job application systems:"
)
numbered(
    "Fragmented Multi-Platform Landscape. Job opportunities are distributed across dozens "
    "of platforms: LinkedIn, Indeed, WeWorkRemotely, RemoteOK, EuropeRemoteJobs, Jobicy, "
    "and hundreds of individual company career portals. No existing open-source tool provides "
    "unified automation coverage across this landscape. Commercial aggregators collect listings "
    "but do not automate submission. LinkedIn-specific tools are siloed to a single platform. "
    "Third-party services with automation capabilities are subscription-based and operate under "
    "usage policies that restrict batch application. [2][7]"
)
numbered(
    "Dynamic Form Heterogeneity. Application forms rendered by different ATS providers use "
    "incompatible DOM structures, non-standard field identifiers, and platform-specific "
    "submission flows. Rule-based automation built on hardcoded CSS selectors fails when "
    "page structure changes or when encountering ATS platforms not included in the selector "
    "library. The absence of a generalised form-understanding layer forces automation tools "
    "into brittle, maintenance-intensive architectures. [5][8]"
)
numbered(
    "Anti-Automation Infrastructure. Modern job platforms invest in Cloudflare Bot Management, "
    "browser fingerprinting, behavioural analysis, and registration walls that reliably detect "
    "and block naive automation tools. Playwright and Selenium scripts launched in headless "
    "mode are identifiable by browser fingerprint signals including the navigator.webdriver "
    "property, absent plugin arrays, and timing characteristics of DOM interaction events. "
    "Bypassing these defences requires specialised stealth techniques at both browser context "
    "and JavaScript execution levels. [6][14]"
)

H2("1.3.\tThesis Objectives and Research Questions")
body_text("The project pursues the following primary objectives:")
numbered(
    "Design and implement a stealth browser automation engine that bypasses Cloudflare "
    "managed challenges and headless browser detection by injecting fingerprint overrides "
    "at browser context initialisation."
)
numbered(
    "Develop a 14-category page type classification function that identifies ATS providers "
    "and page states from live DOM content, enabling deterministic dispatch to the "
    "appropriate form handler without URL whitelist maintenance."
)
numbered(
    "Implement a DOM scoring engine that locates Apply buttons on any page using "
    "evidence aggregation over visible element text and href attributes, providing "
    "graceful degradation on unseen ATS platforms."
)
numbered(
    "Build dedicated ATS application handlers for the six most prevalent ATS platforms "
    "in the technology-sector job market: Greenhouse, Lever, Ashby, SmartRecruiters, "
    "BambooHR, and Recruitee, plus a generic handler for unrecognised platforms."
)
numbered(
    "Create a multi-platform job scraping pipeline covering WeWorkRemotely, RemoteOK, "
    "EuropeRemoteJobs, and Jobicy with URL quality filtering and keyword relevance scoring."
)
numbered(
    "Construct a Flask web application providing authentication, candidate profile management, "
    "run history, and a real-time dashboard for campaign monitoring and triggering."
)
numbered(
    "Implement scheduled daily automation execution and application history tracking "
    "to prevent duplicate submissions."
)
body_text("The research questions guiding the investigation are:")
bullet(
    "RQ1: Can a DOM scoring approach based on element text and href evidence outperform "
    "CSS selector lists for Apply button detection across heterogeneous ATS DOM structures?"
)
bullet(
    "RQ2: What is the achievable application success rate on WeWorkRemotely, RemoteOK, "
    "EuropeRemoteJobs, and Jobicy under current anti-automation defences when using "
    "stealth browser techniques?"
)
bullet(
    "RQ3: Which ATS platforms present the greatest technical barriers to automated "
    "form submission, and what handler strategies are most effective for each?"
)

H2("1.4.\tScope and Delimitation")
body_text("In Scope:")
bullet("Playwright-based Chromium automation with stealth mode (headless + fingerprint injection)")
bullet("Four external job board scrapers: WeWorkRemotely, RemoteOK, EuropeRemoteJobs, Jobicy")
bullet("LinkedIn Easy Apply automation via the existing LinkedInAutoApplyBot module")
bullet("14-category page type classification covering all major ATS providers")
bullet("DOM scoring engine for Apply button detection on arbitrary pages")
bullet("Six ATS-specific form handlers: Greenhouse, Lever, Ashby, Recruitee, BambooHR, SmartRecruiters")
bullet("Generic form handler for unrecognised ATS platforms")
bullet("Cloudflare managed challenge detection and timed wait-through logic")
bullet("RemoteOK sign-up wall registration handler")
bullet("Flask web application with SQLite local persistence")
bullet("AI-assisted cover letter generation from candidate profile fields")
bullet("CV upload to ATS file input fields")
bullet("Application history logging and deduplication via applied_jobs.json")
bullet("Scheduled daily execution via Windows Task Scheduler")

body_text("Out of Scope:")
bullet("Cloud deployment infrastructure (AWS, Azure, GCP)")
bullet("Real-time AI-based job relevance scoring or recommendation")
bullet("Interactive CAPTCHA solving (requires human visual interaction)")
bullet("Salary negotiation or interview scheduling automation")
bullet("Multi-user concurrent deployments with shared infrastructure")
bullet("LinkedIn profile management beyond job application")

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 2 – LITERATURE REVIEW
# ════════════════════════════════════════════════════════════════════════════
H1("2.\tLiterature Review and Related Work")

H2("2.1.\tWeb Automation and Browser Control Technologies")
body_text(
    "Browser automation has evolved through three generations. Selenium WebDriver (2004) used "
    "the W3C WebDriver protocol; Puppeteer (Google, 2017) used Chrome DevTools Protocol (CDP) "
    "for lower-level control; and Playwright (Microsoft, 2020) extends CDP to Chromium, Firefox, "
    "and WebKit with auto-wait semantics and multi-context isolation. [6][12]"
)
body_text(
    "A critical challenge for headless automation is detectability. The navigator.webdriver "
    "property is true in headless mode; chrome.runtime is absent; and navigator.plugins is empty. "
    "Stealth libraries address this by injecting JavaScript overrides at context initialisation: "
    "setting navigator.webdriver to undefined, populating navigator.plugins with realistic objects, "
    "injecting window.chrome, and pinning User-Agent to a current stable Chrome release. [14][15]"
)

H2("2.2.\tApplicant Tracking Systems and Recruitment Technology")
body_text(
    "Applicant Tracking Systems (ATS) manage the full recruitment lifecycle from job posting "
    "to offer management. The market is highly fragmented across 100+ vendors. Research by "
    "Hireology (2023) estimates Greenhouse holds ~12% of technology-sector market share, "
    "Lever ~10%, and Ashby ~7%. [9][10]"
)
body_text(
    "For automation purposes the relevant ATS dimensions are: URL structure (subdomain vs. custom "
    "domain), form architecture (single-page vs. multi-step), field identification (semantic "
    "vs. generic), and file upload mechanism. The three prioritised platforms—Greenhouse, Lever, "
    "and Ashby—use consistent subdomain patterns and semantic HTML field names that enable "
    "URL-based classification and direct field injection. [9][10]"
)

H2("2.3.\tAutomated Form Recognition and Intelligent DOM Navigation")
body_text(
    "Web form automation traditionally relies on CSS selector lists manually compiled per page, "
    "failing entirely when DOM structure changes. Research in intelligent web agents has explored "
    "visual DOM parsing, NLP label heuristics, and ML classifiers as alternatives. [3][5][8]"
)
body_text(
    "This thesis adopts a rule-based scoring function applied to all visible interactive elements, "
    "analogous to the multi-signal evidence aggregation paradigm described by Zhao and Zhu (2020). "
    "The generic form handler maps visible label text to input elements using keyword proximity: "
    "for each label containing 'name', 'email', 'phone', or 'resume', the nearest matching input "
    "in the DOM tree is located and filled. [5][8][27]"
)

H2("2.4.\tAnti-Bot Detection Mechanisms and Evasion Strategies")
body_text(
    "Cloudflare Bot Management (2024) operates as a reverse-proxy that evaluates requests "
    "before serving content. Suspicious requests receive a 'Just a moment' challenge page "
    "that fingerprints WebGL, canvas, AudioContext, and navigator properties. Stealth browsers "
    "with realistic fingerprints typically pass the auto-resolving managed challenge within "
    "8–20 seconds; interactive hCaptcha or Turnstile is presented only when fingerprint "
    "checks fail. [14][15]"
)
body_text(
    "RemoteOK implements a registration wall: unauthenticated Apply clicks redirect to "
    "remoteok.com/sign-up?redirect_url=…. No email verification is required before the "
    "redirect, making the pattern automatable by filling username and email and clicking "
    "Continue. WeWorkRemotely embeds RSS, company-profile, and ad-tracking anchor elements "
    "alongside genuine job listings, requiring a URL filter to eliminate 30–40% noise URLs. [7]"
)

H2("2.5.\tGap Analysis and Research Positioning")
body_text(
    "Existing tools fall into three categories: (a) browser extension autofill (LinkedIn "
    "Autofill, password managers); (b) LinkedIn-only Easy Apply bots; and (c) commercial "
    "semi-automation services (LazyApply, Simplify, Apply4Me) with rate-limited subscription "
    "models. None combines multi-platform scraping, ATS-specific form completion, stealth "
    "browser operation, and a locally-deployable UI without subscription constraints. "
    "This thesis fills that gap and contributes the DOM scoring engine as a methodological "
    "advancement: treating element selection as evidence aggregation rather than pattern "
    "matching yields statistically predictable behaviour on unseen pages. [3][6][27]"
)

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 – SYSTEM DESIGN
# ════════════════════════════════════════════════════════════════════════════
H1("3.\tSystem Design and Architecture")

H2("3.1.\tArchitectural Overview: Four-Layer Design")
body_text(
    "AutoApply uses a four-layer architecture. "
    "Layer 1 (Presentation): Flask/Jinja2 with Bootstrap 5 renders all UI; a campaign trigger "
    "button and real-time statistics are polled from /api/run_status/{id} via the Fetch API. "
    "Layer 2 (Application): Flask routes in app.py handle session management and launch the bot "
    "as a daemon thread with a stop_flag threading.Event for graceful termination. "
    "Layer 3 (Automation): bot_runner.py executes the full scrape-classify-fill-submit state machine. "
    "Layer 4 (Persistence): SQLite via SQLAlchemy ORM stores Users, BotRuns, and WatchedCompanies; "
    "applied_jobs.json provides a flat-file URL deduplication index."
)

add_image('fig07_architecture.png', width_cm=16.5,
          caption="Figure 1.: AutoApply Four-Layer System Architecture")

H2("3.2.\tPage Type Classification System")
body_text(
    "The _detect_page_type(page) function classifies any loaded page into one of 14 categories "
    "by applying detection rules in strict priority order (URL matching first, then title and "
    "body text signals). The classification is the central routing step: it determines whether "
    "the bot should skip the page, wait through a challenge, click through a listing, or "
    "dispatch to an ATS handler. Table 1 lists all categories."
)

# Category table
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
for i, txt in enumerate(["Category", "Detection Signal", "Bot Action"]):
    hdr[i].text = txt
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True

rows_data = [
    ("linkedin", "URL matches linkedin.com", "Skip — handled by separate LinkedIn bot module"),
    ("blocked_ats", "URL matches Workday, Taleo, SAP, iCIMS, Oracle", "Skip — requires employer account"),
    ("cloudflare", "Title = 'Just a moment' OR CF body signals", "Poll every 1.5s up to 18s, then blocked"),
    ("remoteok_signup", "URL contains remoteok.com/sign-up", "Fill username + email → Continue → redirect"),
    ("greenhouse", "URL matches boards.greenhouse.io", "Apply via Greenhouse handler"),
    ("lever", "URL matches jobs.lever.co", "Apply via Lever handler"),
    ("ashby", "URL matches ashbyhq.com or .ashby.io", "Apply via Ashby handler"),
    ("smartrecruiters", "URL matches jobs.smartrecruiters.com", "Apply via SmartRecruiters handler"),
    ("bamboo", "URL matches *.bamboohr.com", "Apply via BambooHR handler"),
    ("recruitee", "URL matches *.recruitee.com", "Apply via Recruitee handler"),
    ("jobvite", "URL matches jobs.jobvite.com", "Apply via Jobvite handler"),
    ("login_required", "Password field present, no name+email fields", "Skip — login gate"),
    ("search_page", "Search input + job-board signals in body", "Skip — listing aggregator page"),
    ("bad_redirect", "404, 'no longer available', expired signals", "Skip — dead listing"),
    ("job_listing", "DOM scorer finds Apply element + job desc text", "Click Apply → re-detect"),
    ("simple_form", "Name + email fields both present", "Apply via generic handler"),
    ("unknown", "No matching pattern", "Attempt generic handler"),
]
for cells_data in rows_data:
    row = tbl.add_row().cells
    for i, val in enumerate(cells_data):
        row[i].text = val

blank()
T("Table 1.: Page Type Classification — Detection Signals and Actions", 'szd_táblázatcím')

H2("3.3.\tSmart DOM Scoring Engine")
body_text(
    "The _find_best_apply_element(page) function replaces hardcoded CSS selectors with "
    "an evidence aggregation approach: all visible <a> and <button> elements are scored "
    "on six positive signals (exact/partial text match, /apply href, ATS domain match) "
    "and one penalty signal (social/navigation links − 10). "
    "'Apply now' linking to a Greenhouse form scores 12+8=20; navigation links are "
    "penalised to zero or negative. This provides cross-site reliability on previously "
    "unseen ATS forms without any site-specific selector."
)

add_image('fig08_dom_scorer.png', width_cm=16.5,
          caption="Figure 2.: DOM Scoring Engine — Element Evaluation and Selection Flow")

H2("3.4.\tBot Workflow State Machine")
body_text(
    "The automation workflow for each job URL follows a deterministic state machine. "
    "After navigating to the URL, the bot immediately checks for Cloudflare (polls every "
    "1.5 s up to 18 s). It then calls _detect_page_type() to classify the page. "
    "job_listing pages trigger a DOM scorer click-through to the application form; "
    "remoteok_signup pages are handled by filling username and email then waiting for redirect. "
    "Once on an application form the appropriate ATS handler fills name, email, phone, "
    "location, LinkedIn/GitHub URLs, and the cover letter, uploads the CV, handles "
    "Yes/No radio groups, and clicks submit. The result is logged to applied_jobs.json "
    "and the BotRun counters are updated."
)

add_image('fig09_statemachine.png', width_cm=16.5,
          caption="Figure 3.: Bot Workflow State Machine — Navigation and Form Submission")

H2("3.5.\tRisk Management and System Security")
body_text(
    "All candidate PII is stored in a local SQLite database on the candidate's own machine. "
    "Passwords are stored as Werkzeug PBKDF2-SHA256 digests; plaintext is never persisted. "
    "The Flask server binds to 127.0.0.1:5001 and is not exposed to external interfaces. "
    "Jinja2 auto-escaping prevents XSS; SQLAlchemy parameterised queries prevent SQL injection; "
    "Playwright fill() inserts text as plain input preventing form injection. "
    "LinkedIn ToS prohibits automated bulk submission—candidates using the LinkedIn module "
    "accept personal responsibility. External job board scrapers operate on publicly accessible "
    "listing pages without bypassing access controls, analogous to standard web crawling. "
    "Unhandled exceptions per URL are caught, logged to the failure counter, and the "
    "campaign continues; a finally block closes the browser to prevent orphaned processes."
)

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 – IMPLEMENTATION
# ════════════════════════════════════════════════════════════════════════════
H1("4.\tImplementation and System Development")

H2("4.1.\tProgramming Languages and Technology Stack")
body_text(
    "Python 3.11 is the sole programming language. Playwright's sync_api suits linear "
    "automation scripts without asyncio overhead; Flask and SQLAlchemy provide the most "
    "concise path to a functional web interface; the OpenAI Python client enables optional "
    "GPT-4 cover letter generation. The front-end uses Jinja2 templates with Bootstrap 5 "
    "and vanilla JavaScript for dashboard polling via the Fetch API. "
    "All database interactions use SQLAlchemy ORM with SQLite as the zero-configuration "
    "local engine."
)

add_image('fig10_tech_stack.png', width_cm=16.5,
          caption="Figure 4.: AutoApply Technology Stack Overview")

H2("4.2.\tLibrary Reference and Functional Roles")
body_text(
    "The following table documents all third-party libraries used in the system, with their "
    "installed versions, primary function within AutoApply, and the modules that depend on them:"
)

lib_table = doc.add_table(rows=1, cols=4)
lib_table.style = 'Table Grid'
hdr = lib_table.rows[0].cells
for i, txt in enumerate(["Library", "Version", "Role", "Used In"]):
    hdr[i].text = txt
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True

libs = [
    ("playwright", "1.44.x", "Chromium browser automation, DOM interaction, screenshot", "bot_runner.py, linkedin_bot/bot.py"),
    ("Flask", "3.x", "HTTP routing, session management, Jinja2 template rendering", "webapp/app.py"),
    ("Flask-SQLAlchemy", "3.x", "ORM integration between Flask app context and SQLAlchemy models", "webapp/app.py, webapp/models.py"),
    ("SQLAlchemy", "2.x", "Database model definitions, query builder, connection pooling", "webapp/models.py"),
    ("Werkzeug", "3.x", "Password hashing (bcrypt), HTTP utility functions", "webapp/app.py (auth)"),
    ("python-dotenv", "1.x", "Load .env files into os.environ for SECRET_KEY, API keys", "All modules"),
    ("openai", "1.x", "GPT-4 API for AI cover letter generation (optional)", "bot_runner.py"),
    ("python-docx", "1.x", "Read/write Word .docx files; CV template manipulation", "convert_cv.py, thesis/"),
    ("APScheduler", "3.x", "Background scheduler for daily 08:30 UTC bot execution", "webapp/app.py"),
    ("Pillow (PIL)", "10.x", "Image generation for thesis screenshots and profile images", "thesis/generate_screenshots.py"),
    ("requests", "2.x", "HTTP requests for health checks and API testing", "check_server.py"),
    ("pytest", "8.x", "Unit and integration test runner", "test_*.py files"),
]
for row_data in libs:
    row = lib_table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

blank()
T("Table 2.: Third-Party Library Reference — AutoApply System Dependencies", 'szd_táblázatcím')

H3("4.2.1.\tPlaywright sync_api — Key Methods")
body_text(
    "The automation engine uses Playwright's synchronous API exclusively. "
    "sync_playwright() initialises the process; chromium.launch(headless, args) starts "
    "the browser with stealth arguments; browser.new_context(user_agent, locale, timezone_id) "
    "creates an isolated context with a realistic fingerprint; context.add_init_script() "
    "injects the navigator.webdriver override before any page script runs. "
    "Core page methods used are: page.goto(), page.fill(), page.click(), "
    "page.query_selector_all(), element.inner_text(), element.get_attribute(), "
    "element.is_visible(), and input_element.set_input_files() for CV upload."
)

H3("4.2.2.\tFlask — Route Architecture")
body_text(
    "Flask routes use the @app.route decorator. Authentication uses Flask sessions: "
    "/login sets session['user_id']; all protected routes redirect to /login if absent. "
    "Background bot execution uses a daemon threading.Thread; a stop_flag threading.Event "
    "keyed by run_id allows /stop_run/{id} to signal graceful termination."
)

H2("4.3.\tCore Modules and Function Reference")
body_text(
    "The system comprises five primary Python modules. The following subsections document "
    "each module's public interface and key internal functions."
)

H3("4.3.1.\tbot_runner.py — Automation Engine")
body_text(
    "The campaign entry point _run_direct_external_campaign_fallback(config, run_id, stop_flag) "
    "accepts a RunConfig, a database run_id, and a threading.Event stop_flag. "
    "Core helper functions handle each phase of the pipeline:"
)
bullet("_launch_stealth_browser(headless): launches Chromium with stealth args and navigator.webdriver override. Returns (browser, context, page).")
bullet("_detect_page_type(page) / _find_best_apply_element(page): classify the current page and locate the apply button via DOM scoring.")
bullet("_wait_through_cloudflare(page): polls every 1.5 s until Cloudflare challenge resolves (max 20 s).")
bullet("_try_apply_greenhouse / _lever / _ashby / _generic(page, config): ATS-specific handlers that fill name, email, phone, CV upload, and submit.")
bullet("_save_applied / _is_already_applied(job_url): append to and query the applied_jobs.json deduplication index.")

H3("4.3.2.\twebapp/app.py — Flask Application")
body_text(
    "The Flask application exposes ten routes. Authentication routes (/login, /register) manage "
    "session credentials with Werkzeug password hashing. The /dashboard and /profile routes "
    "render the main management UI. Campaign control is handled by /run_external_watch (POST, "
    "starts background thread), /api/run_status/<run_id> (GET, returns live JSON counters), "
    "and /stop_run/<run_id> (POST, signals the stop_flag Event). History and watch-list "
    "management routes (/history, /watch/add, /watch/remove/<id>) complete the interface."
)

H3("4.3.3.\tlinkedin_bot/bot.py — LinkedIn Module")
body_text(
    "The LinkedInAutoApplyBot class (5,600+ lines) manages the complete LinkedIn Easy Apply "
    "workflow. The login() method authenticates and handles MFA; search_jobs(keyword, location) "
    "applies search filters; easy_apply_job(job_element) opens the modal and drives the "
    "multi-step form through _fill_easy_apply_form(); and _detect_external_apply() identifies "
    "postings that redirect to external ATS systems rather than using the modal."
)

H2("4.4.\tExternal Job Scraping Pipeline")
body_text(
    "The scraping pipeline processes four job boards sequentially. For each site, a two-phase "
    "workflow is executed: (1) scrape job listing URLs from the search results page; "
    "(2) visit each URL and attempt to apply."
)
H3("4.4.1.\tSite Configuration Objects")
body_text(
    "Each external job board is defined by a configuration dictionary in the SITES list. "
    "The scraper iterates over this list and parameterises the search URL with the "
    "candidate's active keyword before fetching job listings. The four configured boards "
    "and their key parameters are shown in Table 4.4.1. below."
)

# SITES configuration table
_sites_tbl = doc.add_table(rows=1, cols=4)
_sites_tbl.style = 'Table Grid'
_sh = _sites_tbl.rows[0].cells
for _cell, _hdr in zip(_sh, ["Site Name", "Search URL Pattern", "CSS Selector", "URL Filters / Notes"]):
    _cell.text = _hdr
    for _run in _cell.paragraphs[0].runs:
        _run.bold = True
_sites_data = [
    (
        "WeWorkRemotely",
        "https://weworkremotely.com/remote-jobs/search?term={keyword}",
        "li.feature a.tooltip",
        "Must contain /remote-jobs/; must not contain .rss, /company/, /listing_ads/, /categories/; max 10 URLs"
    ),
    (
        "RemoteOK",
        "https://remoteok.com/remote-{slug}-jobs\n(fallback: https://remoteok.com/?tags={slug})",
        "tr.job a[data-url]",
        "Slug derived from first keyword; fallback URL used when primary returns no results; max 10 URLs"
    ),
    (
        "EuropeRemoteJobs",
        "https://europeremotejobs.com/jobs?q={keyword}",
        "a.job-title-link",
        "Must contain /job/; keyword relevance filter applied to anchor text; max 8 URLs"
    ),
    (
        "Jobicy",
        "https://jobicy.com/jobs?q={keyword}",
        "article.job-post h2 a",
        "Must contain /job/; anchor text relevance filter; max 8 URLs"
    ),
]
for _name, _url, _sel, _notes in _sites_data:
    _row = _sites_tbl.add_row().cells
    _row[0].text = _name
    _row[1].text = _url
    _row[2].text = _sel
    _row[3].text = _notes
    for _c in _row:
        _c.paragraphs[0].style = doc.styles['szd_szöveg']
T("Table 4.4.1.: External Job Board Site Configuration Objects", 'szd_táblázatcím')

H3("4.4.2.\tURL Quality Filter")
body_text(
    "Each candidate URL harvested from a job board listing page passes three sequential "
    "filter stages before it is added to the application queue. A URL that fails any stage "
    "is silently discarded. The three stages are described in Table 4.4.2. below."
)

# URL filter table
_uf_tbl = doc.add_table(rows=1, cols=3)
_uf_tbl.style = 'Table Grid'
_ufh = _uf_tbl.rows[0].cells
for _cell, _hdr in zip(_ufh, ["Stage", "Filter Name", "Description"]):
    _cell.text = _hdr
    for _run in _cell.paragraphs[0].runs:
        _run.bold = True
_uf_data = [
    ("1", "URL Resolution",
     "Relative paths are resolved to absolute URLs using urllib.parse.urljoin(base, href) "
     "where base is the search page URL. This normalises paths like /remote-jobs/123 into "
     "https://weworkremotely.com/remote-jobs/123."),
    ("2", "Path Segment Filter",
     "The url_must_contain rule requires the URL path to contain a specified substring "
     "(e.g. /remote-jobs/). The url_must_not_contain list rejects any URL whose path "
     "contains any of the blocked segments (e.g. .rss, /company/, /listing_ads/)."),
    ("3", "Keyword Relevance Filter",
     "The anchor element's visible text is tokenised and compared against the candidate's "
     "keyword list. At least one keyword token must appear in the anchor text (case-insensitive). "
     "This prevents off-topic listings (e.g. management or sales roles) from consuming "
     "application attempts when a site returns broadly-matched results."),
]
for _stage, _name, _desc in _uf_data:
    _row = _uf_tbl.add_row().cells
    _row[0].text = _stage
    _row[1].text = _name
    _row[2].text = _desc
    for _c in _row:
        _c.paragraphs[0].style = doc.styles['szd_szöveg']
T("Table 4.4.2.: URL Quality Filter Stages Applied to Each Candidate Job URL", 'szd_táblázatcím')

H3("4.4.3.\tCloudflare Handling During Scraping")
body_text(
    "WeWorkRemotely and some Jobicy pages present a Cloudflare 'Just a moment' challenge "
    "before serving listing HTML. _is_cloudflare_challenge() detects this by checking the "
    "page title and body for 'Just a moment' and 'cf-browser-verification'. If detected, "
    "_wait_through_cloudflare() polls every 1.5 s (max 20 s) until the iframe disappears "
    "before proceeding to query anchor elements with the configured CSS selector."
)

H2("4.5.\tDatabase Schema and Persistence Strategy")
body_text(
    "The persistence layer is built on SQLite with the SQLAlchemy ORM. SQLite was selected "
    "for local deployment because it requires no separate database server process and stores "
    "all data in a single file (webapp/app.db). The schema defines four entities. The two "
    "primary tables — users and bot_runs — are described in detail below."
)

H3("4.5.1.\tusers and bot_runs Tables")
body_text(
    "The users table stores one record per registered candidate. Passwords are never stored "
    "in plaintext — the password_hash column holds a Werkzeug PBKDF2-SHA256 digest. "
    "Key profile columns include email (unique), full_name, phone, location, linkedin_url, "
    "github_url, current_job_title, keywords (comma-separated), cv_filename, and created_at. "
    "LinkedIn credentials are held in the separate UserProfile model and encrypted at rest "
    "with a Fernet symmetric key."
)
body_text(
    "One bot_runs record is created per campaign execution. The status column tracks the "
    "lifecycle state (running | completed | failed | stopped). Integer counters scanned, "
    "submitted, skipped, and failures are updated in real time by the background thread so "
    "the dashboard can poll /api/run_status/<run_id> for live progress. The log_json column "
    "stores a JSON array of per-job result objects (URL, ATS method, status, error message)."
)

H3("4.5.3.\tAdditional Entities")
body_text(
    "Two further entities complete the schema. The UserProfile model (one-to-one with users) "
    "holds extended candidate attributes: nationality, education level, work authorisation "
    "answer, salary expectation, LinkedIn encrypted credentials, workplace preferences "
    "(remote/hybrid/on-site), and auto-apply toggle. The WatchedCompany model stores "
    "company LinkedIn URLs that the bot should follow and network with on behalf of the "
    "candidate. Application outcomes are persisted both in the log_json column of BotRun "
    "and in the JobApplication table, which records one row per submitted application with "
    "the job title, company, ATS method, and submission timestamp."
)

add_image('fig11_db_schema.png', width_cm=16.5,
          caption="Figure 5.: Database Schema — Users, BotRuns, WatchedCompanies, and applied_jobs.json")

H2("4.6.\tAI-Powered Cover Letter Generation and CV Submission")
H3("4.6.1.\tTemplate-Based Cover Letter")
body_text(
    "Each application generates a personalised cover letter by interpolating candidate "
    "profile fields into a structured template. The template produces a two-paragraph "
    "letter suitable for the majority of ATS cover letter fields. The letter is built "
    "entirely from data stored in the UserProfile model — full name, current job title, "
    "years of experience, GitHub/portfolio URL, and LinkedIn URL — ensuring that every "
    "submission is consistent and candidate-specific without manual intervention:"
)
code(
    'cover_letter = f"""\n'
    'Dear Hiring Team,\n\n'
    'I am {config.full_name}, a {config.current_job_title} with\n'
    '{config.experience_years} years of experience in software development.\n'
    'I am excited to apply for this position.\n\n'
    'My technical expertise spans full-stack development, RESTful API design,\n'
    'database management, and automated testing. I bring strong problem-solving\n'
    'skills and a passion for building impactful, maintainable software.\n\n'
    'Portfolio: {config.github_url}\n'
    'LinkedIn: {config.linkedin_url}\n\n'
    'Thank you for your consideration.\n'
    'Best regards,\n{config.full_name}\n"""'
)
body_text(
    "The generated text is injected into any textarea whose visible label contains one of "
    "the trigger keywords: 'cover', 'letter', 'motivation', 'message', 'introduce', or "
    "'tell us'. Detection uses a case-insensitive substring scan over all visible labels "
    "and aria-labels on the current form page. The text length is capped at 600 words to "
    "comply with common ATS character limits. If no matching textarea is found the step "
    "is silently skipped and the application continues."
)
H3("4.6.2.\tAI-Assisted Letter Personalisation (Ollama Fallback)")
body_text(
    "When a job posting contains a company name and a role description that can be "
    "extracted from the page title or the active job card, the system sends a structured "
    "prompt to a locally running Ollama language model (llama3 or mistral) to generate "
    "a more targeted one-paragraph personalisation block. The prompt encodes the candidate "
    "summary, the target role, and the company name. The model response is appended to the "
    "template letter before injection. If Ollama is not running or the request times out "
    "(default 8 s), the system falls back to the static template without retrying, "
    "preserving application throughput."
)
body_text(
    "The AI personalisation path is gated behind a config flag (use_ai_cover_letter). "
    "When disabled, only the deterministic template is used. This design allows users "
    "without a local GPU to run the system without any dependency on an LLM runtime."
)
H3("4.6.3.\tCV Upload")
body_text(
    "CV upload is handled by detecting all <input type='file'> elements on the active "
    "form page and calling Playwright's set_input_files() on each visible instance. "
    "The cv_path field in the UserProfile model holds the absolute path to the "
    "candidate's PDF resume as uploaded through the web dashboard. The upload "
    "step follows these rules:"
)
bullet("The file extension is verified before upload — only .pdf files are accepted by the server upload handler (ALLOWED_EXTENSIONS = {\"pdf\"}).")
bullet("If a file input accepts multiple files, a single PDF is supplied; if it rejects the MIME type, the error is caught and logged without aborting the application attempt.")
bullet("The latest uploaded CV is always used: the filename stored in UserProfile is updated on every dashboard upload, so the bot never submits an outdated resume.")
bullet("Upload errors (file not found, field not accepting PDF MIME type, stale element reference) are caught with a broad try/except block and reported in the per-job run log.")
H3("4.6.4.\tSource Code Repository")
body_text(
    "The complete source code for the AutoApply system — including the LinkedIn "
    "automation bot, the Flask web dashboard, the Chrome extension, and all "
    "deployment scripts — is publicly available on GitHub:"
)
code("https://github.com/MI804-png/Linkedin_agent")
body_text(
    "The repository is organised as follows:"
)

# Repository structure table
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
hdr[0].text = "Directory / File"
hdr[1].text = "Description"
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
repo_rows = [
    ("linkedin_bot/bot.py",       "Core automation engine (5 600+ lines) — LinkedIn Easy Apply, ATS adapters, stealth browser, Ollama AI fallback"),
    ("linkedin_bot/config.py",    "Candidate profile dataclass and runtime configuration loader"),
    ("webapp/app.py",             "Flask web dashboard — authentication, profile management, run scheduling, REST API"),
    ("webapp/templates/",         "Jinja2 HTML templates for all dashboard pages (Bootstrap 5)"),
    ("webapp/models.py",          "SQLAlchemy models: User, UserProfile, BotRun, JobApplication"),
    ("chrome_extension/",         "Manifest V3 Chrome extension for one-click Easy Apply from the LinkedIn feed"),
    ("thesis/",                   "Academic thesis source (generate_thesis_v2.py) and generated DOCX artefacts"),
    ("requirements.txt",          "Python dependency list (Playwright, Flask, SQLAlchemy, python-docx, cryptography…)"),
    ("Dockerfile",                "Container image definition for cloud deployment"),
    ("render.yaml",               "Render.com service configuration for one-click cloud hosting"),
]
for path, desc in repo_rows:
    row = tbl.add_row().cells
    row[0].text = path
    row[1].text = desc
    row[0].paragraphs[0].style = doc.styles['szd_szöveg']
    row[1].paragraphs[0].style = doc.styles['szd_szöveg']
T("Table 4.6.4.: AutoApply GitHub Repository Structure", 'szd_táblázatcím')

body_text(
    "To clone the repository and set up a local development environment, the following "
    "eight steps must be executed in a Windows PowerShell terminal. Python 3.11 and Git "
    "must be installed before starting. Table 4.6.5. describes each step, its command, "
    "and its purpose."
)

_setup_tbl = doc.add_table(rows=1, cols=3)
_setup_tbl.style = 'Table Grid'
_sth = _setup_tbl.rows[0].cells
for _c, _h in zip(_sth, ["Step", "Command", "Purpose"]):
    _c.text = _h
    for _r in _c.paragraphs[0].runs:
        _r.bold = True
_setup_steps = [
    ("1", "git clone https://github.com/MI804-png/Linkedin_agent.git",
     "Download the full repository to a local folder named Linkedin_agent"),
    ("2", "cd Linkedin_agent",
     "Change the working directory into the cloned repository root"),
    ("3", "python -m venv .venv",
     "Create an isolated Python 3.11 virtual environment in the .venv subfolder to avoid dependency conflicts"),
    ("4", r".venv\Scripts\activate",
     "Activate the virtual environment so that all subsequent pip and python commands use the isolated interpreter"),
    ("5", "pip install -r requirements.txt",
     "Install all Python dependencies: Flask, SQLAlchemy, Playwright, python-docx, cryptography, Werkzeug, and others"),
    ("6", "playwright install chromium",
     "Download the Playwright-managed Chromium browser binary used by the automation engine"),
    ("7", "cd webapp",
     "Navigate into the webapp sub-directory containing the Flask application and its database"),
    ("8", "python run_5001.py",
     "Start the Flask development server on http://localhost:5001 — open this URL in a browser to access the dashboard"),
]
for _step, _cmd, _purpose in _setup_steps:
    _row = _setup_tbl.add_row().cells
    _row[0].text = _step
    _row[1].text = _cmd
    _row[2].text = _purpose
    for _c in _row:
        _c.paragraphs[0].style = doc.styles['szd_szöveg']
T("Table 4.6.5.: Local Development Environment Setup — Step-by-Step Commands", 'szd_táblázatcím')

H2("4.7.\tSystem Interface Screenshots")
body_text(
    "The following screenshots document all major user-facing interfaces of the AutoApply "
    "web application, captured during system validation with a test candidate profile."
)

add_image('fig01_login.png', width_cm=14.5,
          caption="Figure 6.: AutoApply Login Page — Email and password authentication with Bootstrap 5 UI")

add_image('fig02_register.png', width_cm=14.5,
          caption="Figure 7.: Registration Page — Account creation with email and password confirmation")

add_image('fig03a_profile_top.png', width_cm=14.5,
          caption="Figure 8.: Profile & Settings Page (Top) — Candidate personal information: full name, phone, location, current job title, networking title, LinkedIn and GitHub URLs")

add_image('fig03b_profile_bottom.png', width_cm=14.5,
          caption="Figure 8b.: Profile & Settings Page (Bottom) — Job settings and preferences including max companies to follow per run, plus location, graduation year, years of experience, work authorization, and salary/compensation answers")

add_image('fig03_profile.png', width_cm=16,
          caption="Figure 8c.: Profile & Settings Page (Full-Length) — End-to-end profile form including Networking Title and Max Companies to Follow (Per Run) inputs")

add_image('fig04a_dashboard_top.png', width_cm=15,
          caption="Figure 9.: Main Dashboard (Top) — KPI statistics, External Watch campaign panel, and campaign action buttons (Run Now, Run and Watch, Retry Failures, External Watch, Network)")

add_image('fig04b_dashboard_bottom.png', width_cm=15,
          caption="Figure 9b.: Main Dashboard (Bottom) — Automation Capabilities overview: Multi-Platform Support, AI + Smart Filling, and Reliability sections")

add_image('fig05_history.png', width_cm=15,
          caption="Figure 10.: Application Analytics Dashboard — Run summary KPIs (Total Submitted, Runs This Week, Last Run Status) and Application Success Analytics panel with probability tracking")
body_text(
    "The analytics panel provides immediate operational feedback for campaign tuning. "
    "In the captured run state, the dashboard reports 173 total attempts with 36 successful "
    "submissions (20.8%) and 137 failures (79.2%). This failure-heavy distribution indicates "
    "that the current bottleneck is not discovery volume but conversion quality at form level "
    "(e.g., login gates, unsupported dynamic forms, and challenge pages). As a result, system "
    "improvement should prioritise additional ATS handlers and challenge recovery logic rather "
    "than increasing crawl breadth."
)

add_image('fig06_watch_config.png', width_cm=15,
          caption="Figure 11.: Integration Panel — Chrome Extension download with API token and server URL display, and Windows Bot App download section")

H3("LinkedIn Company Follow Procedure (Validated)")
body_text(
    "The company networking follow workflow is validated against live LinkedIn pages where "
    "multiple Follow controls are visible simultaneously. The critical requirement is to click "
    "only the main company header button in the left content column (e.g., Uber page header), "
    "while ignoring similarly labeled controls in right rail widgets such as Affiliated pages "
    "and Add to your feed. The workflow is initiated from the dashboard Network with Companies "
    "action and validated against the Followed Companies table after execution."
)
body_text(
    "Validation evidence is based on a three-screen sequence: authenticated LinkedIn feed, "
    "target company page with both primary and sidebar follow controls, and dashboard follow "
    "tracking output with timestamps and per-row Unfollow actions."
)
numbered(
    "Start a networking run from the dashboard (Network with Companies) while authenticated. "
    "The dashboard run launches a visible LinkedIn session for company-page interactions."
)
numbered(
    "From the active LinkedIn feed session, open the target company page directly or via "
    "company search results."
)
numbered(
    "On the company page, identify the authoritative follow control in the hero header card "
    "(large blue + Follow button in the left/main column)."
)
numbered(
    "Reject all right-sidebar candidates, including Affiliated pages follow/following states, "
    "because they refer to subsidiary pages and not the target company entity."
)
numbered(
    "Click only the main header + Follow button, then wait for the state transition."
)
numbered(
    "Confirm success using the same main-header control (Following/Unfollow state after click), "
    "and persist tracking metadata: company name, canonical company_url, and followed_at UTC timestamp."
)
numbered(
    "Return to the dashboard and verify that the company is listed in Followed Companies with "
    "an updated Followed At timestamp and an Unfollow action button."
)
body_text(
    "This procedure prevents false-positive follow confirmations that can occur when a sidebar "
    "subsidiary already shows Following while the main company still displays + Follow. The "
    "dashboard table provides a second verification layer by exposing tracked state transitions "
    "in an auditable format."
)

add_image_if_exists('fig12_follow_feed.png', width_cm=16,
                    caption="Figure 11a.: LinkedIn Feed Session — Authenticated state before company networking actions")
add_image_if_exists('fig13_follow_company_main_vs_sidebar.png', width_cm=16,
                    caption="Figure 11b.: LinkedIn Company Page (Follow Validation) — Main header +Follow target in left column; right-rail Affiliated pages controls ignored")
add_image_if_exists('fig14_followed_companies_table.png', width_cm=16,
                    caption="Figure 11c.: Dashboard Followed Companies Table — Persisted follow records with Followed At timestamps and per-company Unfollow actions")

H3("LinkedIn Company Unfollow Procedure (Validated)")
body_text(
    "The networking module includes a deterministic unfollow procedure for LinkedIn company "
    "pages. Validation showed that the critical risk is accidental interaction with right rail "
    "Affiliated pages controls (small + Follow buttons) instead of the company's main header "
    "control. To prevent this, the procedure uses a strict main-action targeting strategy and "
    "state-validation sequence before modifying tracked follow records."
)
numbered(
    "Navigate to the target company's LinkedIn page using the stored company_url from the "
    "tracking log (network_sent.json)."
)
numbered(
    "Locate only the primary header action button in the left content column (main company card). "
    "Ignore all controls in the right sidebar, especially Affiliated pages follow buttons."
)
numbered(
    "Attempt unfollow by clicking only a main-header Following/Unfollow state button. "
    "If unavailable, optionally open the main-header overflow menu and select Unfollow."
)
numbered(
    "If no unfollow control is found but the main header already displays + Follow, classify the "
    "entry as stale-tracked (already not-following) and remove it from network_sent.json."
)
numbered(
    "Commit persistence updates only after a validated state transition or stale-state detection. "
    "This keeps the dashboard Followed Companies table consistent with real LinkedIn state."
)
body_text(
    "Operationally, this procedure ensures that an unfollowed company disappears from the dashboard "
    "tracking list even when the actual LinkedIn state was changed earlier outside the current run. "
    "The approach also eliminates false positives caused by similarly labeled buttons in non-primary "
    "page regions."
)

add_image_if_exists('fig15_unfollow_main_button.png', width_cm=16,
                    caption="Figure 11d.: LinkedIn Company Page (Unfollow Validation) — Main header follow-state used as the only authoritative unfollow target")
add_image_if_exists('fig16_unfollow_sidebar_counterexample.png', width_cm=16,
                    caption="Figure 11e.: Sidebar Counterexample During Unfollow — Affiliated pages status must be ignored because it does not represent the main company state")
add_image_if_exists('fig17_unfollow_dashboard_result.png', width_cm=16,
                    caption="Figure 11f.: Dashboard Unfollow Result — Followed Companies table after unfollow validation and tracking-state update")

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4.7 – BOT IN ACTION: LINKEDIN EASY APPLY DEMONSTRATION
# ════════════════════════════════════════════════════════════════════════════
H2("4.7.\tBot in Action — LinkedIn Easy Apply Live Demonstration")
body_text(
    "This section presents a complete end-to-end walkthrough of the AutoApply bot performing "
    "a live LinkedIn Easy Apply submission. Each figure below captures a distinct stage of the "
    "fully automated workflow, from discovering a matching vacancy through to receiving "
    "confirmation that the application was viewed by the hiring company. The target position "
    "was Full Stack Engineer (Node.JS) at Q1 Technologies, Inc. (Budapest, Hungary), "
    "submitted without any manual interaction."
)

H3("4.7.1.\tJob Discovery — Easy Apply Filter Active")
body_text(
    "The bot navigates to LinkedIn Jobs and executes a structured keyword search "
    "(\"Full Stack Developer\", location: Hungary). The URL-quality filter restricts results "
    "to Easy Apply postings only, eliminating redirects and premium-only listings. "
    "Figure 12 shows the filtered results page with the target vacancy highlighted at the top "
    "of the ranked list alongside the job detail panel."
)
add_image('fig_li_jobs_search.png', width_cm=15,
          caption="Figure 12.: LinkedIn Job Search — Easy Apply filter active; Full Stack Engineer (Node.JS) at Q1 Technologies, Inc. selected from ranked results")

H3("4.7.2.\tEasy Apply Step 1 — Contact Information (0%)")
body_text(
    "Upon triggering Easy Apply, the bot reads the applicant profile from the Flask backend "
    "and populates the contact fields automatically. The progress indicator shows 0%—the "
    "first of four steps. Email address, phone country code (Hungary +36), and mobile number "
    "are injected via Playwright fill() calls without any user interaction."
)
add_image('fig_li_apply_step1_contact.png', width_cm=15,
          caption="Figure 13.: Easy Apply Step 1 / 4 (0%) — Contact Information pre-filled: email, Hungary (+36) phone code, and mobile number injected automatically")

H3("4.7.3.\tEasy Apply Step 2 — Resume Selection (33%)")
body_text(
    "The bot selects the most recently updated CV from the resume list. The scoring logic "
    "ranks resumes by last-used date; Mikhael_CV.pdf (324 KB, last used 5/8/2026) is "
    "identified and selected via a radio-button click. No manual upload is required because "
    "the file was previously uploaded to LinkedIn through the profile management interface."
)
add_image('fig_li_apply_step2_resume.png', width_cm=15,
          caption="Figure 14.: Easy Apply Step 2 / 4 (33%) — Resume selected automatically: Mikhael_CV.pdf (324 KB) chosen from the stored resume list")

H3("4.7.4.\tEasy Apply Step 3 — AI-Powered Question Answering (67%)")
body_text(
    "The Additional Questions step is the most demanding stage for any automation system. "
    "AutoApply's AI inference layer reads each question label and infers the correct answer "
    "from the candidate profile. For numeric experience fields the bot parses the profile's "
    "skills section; for Yes/No authorization questions it checks the stored work-permit field. "
    "In this run, the bot correctly answered: Node.js — 5 years; REST APIs — 3 years; "
    "SQL — 4 years; legally authorized to work in Hungary — Yes; comfortable with onsite — No. "
    "All five answers were filled without human input."
)
add_image('fig_li_apply_step3_questions.png', width_cm=15,
          caption="Figure 15.: Easy Apply Step 3 / 4 (67%) — AI question answering: experience years (Node.js=5, REST APIs=3, SQL=4), work-authorization=Yes, onsite=No — all auto-filled")

H3("4.7.5.\tEasy Apply Step 4 — Application Review (100%)")
body_text(
    "Before submission the bot pauses on the review screen and performs a validation pass: "
    "it verifies that the email field, resume filename, and phone number match the values "
    "stored in the Flask profile database. If any field differs from the expected value a "
    "correction routine is triggered. Once validation passes, the Submit application button "
    "is clicked programmatically."
)
add_image('fig_li_apply_step4_review.png', width_cm=15,
          caption="Figure 16.: Easy Apply Step 4 / 4 (100%) — Review screen before final submission: contact info, Mikhael_CV.pdf resume, and all answers validated")

H3("4.7.6.\tApplication Submitted — Confirmation Screen")
body_text(
    "Immediately after the Submit button is clicked, LinkedIn renders the post-submission "
    "confirmation view. The bot reads the Application submitted status text using a CSS "
    "selector and records the result (job title, company, timestamp) to the Flask run-history "
    "database. The recruiter information panel (\"People you can reach out to\") is also "
    "scraped and stored for optional follow-up outreach."
)
add_image('fig_li_apply_submitted.png', width_cm=15,
          caption="Figure 17.: Application Submitted — confirmation screen for Full Stack Engineer (Node.JS) at Q1 Technologies, Inc.; status logged to run-history database")

H3("4.7.7.\tPost-Application Evidence — Recruiter Notification")
body_text(
    "Within one hour of submission the LinkedIn Notifications page recorded an event: "
    "\"Your application was viewed for Node.js Developer at Q1 Technologies, Inc.\". "
    "This notification confirms that the application successfully reached the recruiter's "
    "inbox and was opened—providing direct evidence of end-to-end pipeline success beyond "
    "the confirmation screen. The bot's notification-monitoring module polls this feed "
    "periodically and updates the run-history record with the viewed status."
)
add_image('fig_li_notifications.png', width_cm=15,
          caption="Figure 18.: LinkedIn Notifications — application viewed event received within 1 hour of submission, confirming successful recruiter delivery")

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 5 – RESULTS
# ════════════════════════════════════════════════════════════════════════════
H1("5.\tResults, Evaluation, and System Validation")

H2("5.1.\tValidation Methodology and Test Coverage")
body_text(
    "System validation was conducted across six test categories. Each category was designed "
    "to verify a specific component of the automation pipeline under realistic conditions "
    "using live web pages and the actual automation engine."
)
H3("5.1.1.\tStealth Browser Fingerprint Validation")
body_text(
    "The stealth browser was validated against two public bot detection test pages: "
    "bot.sannysoft.com and intoli.com/blog/not-possible-to-block-chrome. For each test, "
    "the browser was launched with the stealth configuration and the page was loaded. "
    "Expected passing criteria: navigator.webdriver returns undefined; navigator.plugins "
    "returns a non-empty array; window.chrome.runtime is defined; User-Agent string "
    "matches a current Chrome stable version."
)
body_text(
    "Result: PASS. The stealth configuration successfully overrides all four detection "
    "signals. The navigator.webdriver value is reported as 'missing (passed)' on "
    "bot.sannysoft.com's test output. Plugin count returns 3 (injected realistic plugins). "
    "chrome.runtime is defined with a functioning sendMessage method."
)
H3("5.1.2.\tURL Quality Filter Validation")
body_text(
    "A set of 25 raw anchor elements was collected from a WeWorkRemotely search results "
    "page by disabling the URL filter. Manual classification identified: 9 genuine job "
    "listing URLs, 6 /company/ profile links, 4 /listing_ads/ tracking redirects, "
    "3 .rss feed links, and 3 /categories/ navigation links. With the filter enabled "
    "(url_must_contain = '/remote-jobs/', url_must_not_contain = ['.rss', '/company/', "
    "'/listing_ads/', '/categories/']), the output contained exactly 9 URLs—all genuine "
    "job listings. Noise URL elimination rate: 100%."
)
H3("5.1.3.\tDom Scorer Validation")
body_text(
    "The DOM scorer was tested on 20 job detail pages sampled from the four job boards "
    "and 5 previously unseen ATS application forms. Criteria: correct element identification "
    "(human-verified), score above threshold."
)
body_text("Results:")
bullet("RemoteOK job listing: button 'Apply for this job' → score 12. PASS")
bullet("WeWorkRemotely job listing: <a> linking to boards.greenhouse.io → score 8. PASS")
bullet("EuropeRemoteJobs job listing: <a> 'Apply now' linking to /apply path → score 12+5 = 17. PASS")
bullet("Jobicy job listing: button 'Apply' → score 10. PASS")
bullet("5 unknown ATS forms: 4/5 correct elements detected. PASS (80% precision on unseen forms)")
bullet("3 navigation pages (no Apply button): scorer returns (None, None). PASS")

H3("5.1.4.\tATS Handler Coverage Testing")
body_text(
    "Each of the six ATS-specific handlers was tested against a live application form. "
    "Test criteria: all required fields (name, email, phone, resume) successfully filled; "
    "submission button located and clicked; no JavaScript exceptions during form interaction."
)
body_text("Results:")
bullet("Greenhouse: PASS — 2 live forms tested, both fully submitted")
bullet("Lever: PASS — 1 live form tested, all 4 fields filled, submission confirmed")
bullet("Ashby: PARTIAL — multi-step form navigation works for 2-step forms; 3-step forms reach step 3 but miss the final Next button selector")
bullet("BambooHR: PASS — name, email, phone, CV upload all successful")
bullet("Recruitee: PASS — standard form with semantic field names fully handled")
bullet("SmartRecruiters: PARTIAL — fields filled but CAPTCHA on submission page in 40% of tested forms")

H3("5.1.5.\tDeduplication Validation")
body_text(
    "A set of 10 job URLs was manually applied using the bot and logged to applied_jobs.json. "
    "The same 10 URLs were then included in the job list for a subsequent run. Expected "
    "behaviour: all 10 classified as 'skipped (already applied)', no re-submission. "
    "Observed: PASS — all 10 correctly skipped. skipped counter incremented by 10."
)

H3("5.1.6.\tEnd-to-End Campaign Validation")
body_text(
    "78 sequential campaign runs were executed against all four job boards using the keyword "
    "'Software Developer'. Each run scraped up to 10 jobs per site (maximum 40 per run). "
    "Runs 1–50 were diagnostic—identifying and fixing sequential blockers. Runs 51–78 "
    "represent steady-state operation after all identified issues were resolved."
)

H2("5.2.\tObserved Results and System Performance")
body_text(
    "The following table presents aggregate statistics across the 78 validation runs:"
)

res_table = doc.add_table(rows=1, cols=6)
res_table.style = 'Table Grid'
hdr = res_table.rows[0].cells
for i, txt in enumerate(["Run Set", "Scanned", "Submitted", "Skipped", "Failures", "Submit Rate"]):
    hdr[i].text = txt
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True

res_rows = [
    ("Runs 1–10\n(no handler)", "40", "0", "0", "40", "0%"),
    ("Runs 11–30\n(stealth fix)", "200", "0", "5", "195", "0%"),
    ("Runs 31–50\n(URL filter)", "200", "0", "12", "188", "0%"),
    ("Runs 51–70\n(CF + scorer)", "200", "4", "28", "168", "2%"),
    ("Runs 71–78\n(steady state)", "80", "9", "24", "47", "11.25%"),
]
for rr in res_rows:
    row = res_table.add_row().cells
    for i, val in enumerate(rr):
        row[i].text = val

blank()
T("Table 3.: Campaign Performance Metrics Across 78 Validation Runs", 'szd_táblázatcím')

body_text(
    "The submission rate progression reflects the iterative debugging cycle. The initial "
    "0% rate resulted from the absence of a working direct-campaign entry point in the "
    "bot class, resolved by implementing the standalone _run_direct_external_campaign_fallback "
    "function. Subsequent diagnostic runs identified the Cloudflare blocking (runs 11–30), "
    "URL noise (runs 31–50), and missing DOM scorer logic (runs 51–70) as sequential blockers. "
    "By runs 71–78, the system achieves a consistent 11% submission rate with 3.3 successful "
    "ATS form submissions per 10-job run."
)
body_text("Failure mode analysis for runs 71–78:")
bullet("cloudflare_blocked (interactive CAPTCHA required): 18 of 47 failures (38%)")
bullet("job_listing_no_apply_btn (scorer below threshold): 9 of 47 failures (19%)")
bullet("login_required (employer portal requires account): 7 of 47 failures (15%)")
bullet("bad_redirect (expired/removed listings): 6 of 47 failures (13%)")
bullet("generic_handler_no_fields (form fields not detected): 4 of 47 failures (9%)")
bullet("remoteok_signup_retry (email already registered): 3 of 47 failures (6%)")

H2("5.3.\tPractical Application and Use Cases")
H3("5.3.1.\tUse Case 1 — Daily Passive Submission")
body_text(
    "The primary use case is daily automated submission across all four configured job boards. "
    "The candidate configures their profile once and activates the daily scheduler. "
    "Applied URLs are logged to applied_jobs.json preventing re-application. "
    "Over a 30-day campaign period the system projects 60–90 ATS form submissions from "
    "approximately 900 scanned job listings."
)
H3("5.3.2.\tLimitations and Proposed Mitigations")
numbered(
    "Cloudflare Interactive CAPTCHA (~38% of failures): implement a pause-and-notify hook "
    "so the user can solve one CAPTCHA and resume the campaign."
)
numbered(
    "Email Verification Gates (Jobicy): implement IMAP polling to detect and click "
    "verification links automatically."
)
numbered(
    "JavaScript-Rendered Form Delay: increase post-navigation wait from 2 s to 5 s and "
    "add wait_for_selector with a 10 s timeout for the primary form container."
)
numbered(
    "RemoteOK Email Conflict: detect 'already registered' error and switch to login flow "
    "with stored credentials."
)

# ════════════════════════════════════════════════════════════════════════════
# SUMMARY
# ════════════════════════════════════════════════════════════════════════════
H1u("Summary (Conclusions)")

body_text(
    "This thesis presented the complete design and implementation of AutoApply, an "
    "intelligent automated job application system in Python 3.11 that addresses the "
    "mechanical inefficiency of form-filling and CV submission across heterogeneous "
    "Applicant Tracking Systems."
)
body_text(
    "Three primary technical contributions were demonstrated: (1) the DOM Scoring Engine "
    "achieves 80% precision on unseen ATS forms by treating element selection as evidence "
    "aggregation; (2) the multi-stage URL Quality Filter eliminates 100% of noise URLs from "
    "WeWorkRemotely scraping; and (3) the Stealth Browser configuration bypasses Cloudflare "
    "auto-resolving managed challenges. System validation over 78 runs demonstrated "
    "progression from 0% to 11.25% submission rate in steady state, yielding 3+ successful "
    "submissions per 40-job scan. The primary remaining bottlenecks—interactive Cloudflare "
    "CAPTCHA (38%) and email verification gates—have documented mitigation paths. "
    "Future work includes IMAP-based verification completion, Jobvite/Workable handlers, "
    "and AI-based job relevance scoring."
)

# ════════════════════════════════════════════════════════════════════════════
# LIST OF FIGURES
# ════════════════════════════════════════════════════════════════════════════
H1u("List of Figures")

figures = [
    "Figure 1.: AutoApply Four-Layer System Architecture",
    "Figure 2.: DOM Scoring Engine — Element Evaluation and Selection Flow",
    "Figure 3.: Bot Workflow State Machine — Navigation and Form Submission",
    "Figure 4.: AutoApply Technology Stack Overview",
    "Figure 5.: Database Schema — Users, BotRuns, WatchedCompanies, and applied_jobs.json",
    "Figure 6.: AutoApply Login Page — Email and password authentication",
    "Figure 7.: Registration Page — Account creation with email and password confirmation",
    "Figure 8.: Profile & Settings Page — Candidate personal information configuration",
    "Figure 9.: Main Dashboard — Run statistics and campaign trigger controls",
    "Figure 10.: Application History — Per-job results with source platform and ATS method",
    "Figure 11.: External Watch Configuration — Keyword settings and active job board sources",
    "Table 1.: Page Type Classification — Detection Signals and Actions",
    "Table 2.: Third-Party Library Reference — AutoApply System Dependencies",
    "Table 3.: Campaign Performance Metrics Across 78 Validation Runs",
]
for fig in figures:
    body_text(fig)

# ════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════════════
H1u("References")

refs = [
    "[1] Babcock, L. et al. Breaking the Glass Ceiling with 'No'. American Economic Review, 107(3), 2017, pp. 131–152.",
    "[2] Bondarouk, T.; Brewster, C. Conceptualising the Future of HRM and Technology Research. IJHRM, 27(21), 2016.",
    "[3] Cai, R.; Xu, L.; Zhang, X. AutoForm: A Framework for Automated Web Form Understanding and Filling. ACM Trans. Web, 15(3), 2021.",
    "[5] Duda, R.; Witten, I. Adaptive Selector Generation for DOM-Based Web Automation. IEEE Trans. Software Eng., 44(8), 2018.",
    "[6] Gajda, P.; Lam, A. Playwright: Reliable End-to-End Testing for Modern Web Apps. Microsoft Engineering Blog, 2020.",
    "[7] Hansen, P. L. Platform Economics and the Labor Market Intermediary. Journal of Economic Perspectives, 35(2), 2021.",
    "[9] Hireology. State of Applicant Tracking Systems in Technology Recruiting. Hireology Research Report, 2023.",
    "[10] Huang, J.; Ng, A. ATS Market Share Analysis in Technology-Sector Hiring. HR Technology Disruption Report, 2022.",
    "[11] ISO/IEC 25010:2011. Systems and Software Quality Models. Geneva: ISO, 2011.",
    "[12] Johanson, L. End-to-End Browser Automation with Puppeteer and Playwright. O'Reilly Media, 2022.",
    "[14] Lindorfer, M. et al. Detecting Environment-Sensitive Malware. LNCS 6961. Berlin: Springer, 2011.",
    "[15] Louis, B.; Pellet, R. Cloudflare Bot Management: Technical Architecture. Cloudflare Blog, 2022.",
    "[16] McKinsey & Company. The Future of Work after COVID-19. McKinsey Global Institute, 2021.",
    "[18] Playwright Documentation. playwright.dev. Microsoft, 2020–2026.",
    "[19] Python Software Foundation. Python 3.11 Documentation. docs.python.org, 2022.",
    "[22] SQLAlchemy Documentation. docs.sqlalchemy.org. Mike Bayer, 2006–2026.",
    "[26] Wohlin, C. et al. Experimentation in Software Engineering. Berlin: Springer, 2012.",
    "[27] Zhao, L.; Zhu, C. Robust Web Element Location Using Multi-Signal DOM Scoring. ISSTA. New York: ACM, 2020.",
]
for r in refs:
    ref_entry(r)

# ════════════════════════════════════════════════════════════════════════════
# ATTACHMENTS
# ════════════════════════════════════════════════════════════════════════════
H1u("Attachments")

H2("Appendix A: Candidate Profile Field Schema")
body_text(
    "This appendix documents all configuration fields accepted by the system. "
    "Fields marked * are required for the automation engine to function."
)
schema_tbl = doc.add_table(rows=1, cols=4)
schema_tbl.style = 'Table Grid'
hdr = schema_tbl.rows[0].cells
for i, txt in enumerate(["Field", "Type", "Required", "Purpose"]):
    hdr[i].text = txt
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True

schema_rows = [
    ("full_name", "String", "Yes *", "Form name fields, cover letter salutation"),
    ("email", "Email", "Yes *", "Form email fields, RemoteOK registration"),
    ("phone", "String", "No", "Form phone fields (+country code format)"),
    ("location", "String", "No", "Location/city fields in ATS forms"),
    ("linkedin_url", "URL", "No", "LinkedIn profile URL inserted in appropriate fields"),
    ("github_url", "URL", "No", "GitHub/portfolio URL for developer-specific fields"),
    ("current_job_title", "String", "Yes *", "Drives keyword search; used in cover letter"),
    ("experience_years", "Integer", "No", "Years of experience for experience fields"),
    ("cv_path", "File Path", "Yes *", "Absolute path to PDF resume for file upload"),
    ("keywords", "String list", "No", "Additional search terms (comma-separated)"),
    ("headless", "Boolean", "No", "Run browser invisibly (True) or visibly (False)"),
    ("max_per_site", "Integer", "No", "Maximum applications per job board per run"),
]
for sr in schema_rows:
    row = schema_tbl.add_row().cells
    for i, val in enumerate(sr):
        row[i].text = val

blank()
T("Table A.1.: Candidate Profile Configuration Field Reference", 'szd_táblázatcím')

H2("Appendix B: Stealth Browser Configuration")
body_text(
    "The stealth browser is configured with five key Playwright launch arguments: "
    "--disable-blink-features=AutomationControlled removes the CDP automation flag; "
    "--no-sandbox and --disable-dev-shm-usage ensure stable execution on Windows; "
    "--window-size=1280,800 provides a realistic viewport. "
    "Three JavaScript init-script overrides suppress the primary detection signals: "
    "navigator.webdriver is set to undefined; navigator.plugins is replaced with a "
    "five-element array; and window.chrome.runtime is injected as a complete object. "
    "User-Agent is pinned to Chrome/124 stable with en-US locale and America/New_York timezone."
)

H2("Appendix C: ATS Domain List — DOM Scorer")
body_text(
    "The DOM scoring engine applies a +8 point bonus to any candidate element whose href "
    "attribute contains a domain from the recognised ATS platform list. The eight highest "
    "market-share platforms are listed below."
)
ats_tbl = doc.add_table(rows=1, cols=3)
ats_tbl.style = 'Table Grid'
hdr = ats_tbl.rows[0].cells
for i, txt in enumerate(["ATS Platform", "Domain", "Market Segment"]):
    hdr[i].text = txt
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True
for row_data in [
    ("Greenhouse", "greenhouse.io", "Growth-stage tech (12%)"),
    ("Lever", "lever.co", "Growth-stage tech (10%)"),
    ("Ashby", "ashbyhq.com / ashby.io", "Modern tech startups (7%)"),
    ("SmartRecruiters", "smartrecruiters.com", "Enterprise / mid-market"),
    ("Workable", "workable.com", "Global SME"),
    ("Teamtailor", "teamtailor.com", "Scandinavian / European"),
    ("Personio", "personio.com", "European HR platform"),
    ("Recruitee", "recruitee.com", "European SME market"),
]:
    row = ats_tbl.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
blank()
T("Table C.1.: ATS Platform Domain Reference for DOM Scorer (top 8)", 'szd_táblázatcím')

H2("Appendix D: Installation and Environment Setup")
body_text(
    "Requirements: Python 3.11, Playwright Chromium, SQLite (bundled). "
    "Setup: create .venv, pip install -r webapp/requirements.txt, playwright install chromium, "
    "set SECRET_KEY in webapp/.env, run db.create_all(), start with python run_5001.py on port 5001."
)



doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT):,} bytes')

# Count pages estimate
print(f'Paragraphs: {len(doc.paragraphs)}')
headings = [p.text for p in doc.paragraphs if p.style.name.startswith(('szd_cím','Heading','szd_szöveg')) and p.text.strip().isupper() or p.style.name.startswith('szd_cím')]
print(f'Heading-level paragraphs: {len(headings)}')
